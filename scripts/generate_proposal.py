# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"d:\Spain_Hotel_Camera_Management_System\Hotel_VMS_Project_Proposal_IntelliForge.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=RGBColor(0x0B, 0x2E, 0x4A))
    return p


def add_para(doc, text, size=11, bold=False, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=10, center=True)
        shade_cell(table.rows[0].cells[i], "0B2E4A")
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], str(val), size=10)
            if r_idx % 2 == 1:
                shade_cell(table.rows[r_idx + 1].cells[c_idx], "F2F6FA")
    doc.add_paragraph()
    return table


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PROJECT PROPOSAL")
    set_run_font(r, size=22, bold=True, color=RGBColor(0x0B, 0x2E, 0x4A))

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Intelligent Video Surveillance System for Hotel\nMilestone XProtect + AI Analytics")
    set_run_font(r, size=14, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        "Prepared by: IntelliForge\n"
        "Client project: Hotel videovigilancia inteligente (Valencia)\n"
        "Delivery mode: 100% Remote (no on-site visit required)\n"
        "Date: August 2026"
    )
    set_run_font(r, size=11)

    add_para(doc, "")

    # 1. Executive summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This proposal covers the complete remote implementation and configuration of an intelligent "
        "video surveillance workflow for a hotel with 30 Safire Smart E1 Series cameras (already installed) "
        "and a Safire Smart B2 NVR (already operating). The core platform is Milestone XProtect, extended "
        "with AI analytics (BriefCam or equivalent) for facial recognition and multi line-crossing, plus "
        "real-time alerting to mobile and WhatsApp/SMS.",
    )
    add_para(
        doc,
        "A interactive operator-workflow demo has already been prepared to illustrate live monitoring, "
        "watchlists, rules, alerts, search/reporting, and GDPR controls. The production delivery will be "
        "configured on the real Milestone/Safire environment (not a custom replacement VMS).",
    )

    # 2. Project understanding
    add_heading(doc, "2. Project Understanding", 1)
    add_para(doc, "Existing assets:", bold=True, space_after=4)
    for item in [
        "30× Safire Smart E1 Series cameras (Bullet/Turret 6MP) — installed",
        "Safire Smart B2 NVR — in operation",
        "Target platform: Milestone XProtect + AI analytics module",
        "Remote delivery preferred; no Valencia on-site visit required",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    add_para(doc, "", space_after=4)
    add_para(doc, "Required functional outcomes:", bold=True, space_after=4)
    for item in [
        "Facial recognition with employee / guest / blocked watchlists",
        "Multiple virtual line-crossing rules (room doors + common areas)",
        "Real-time alerts to mobile and WhatsApp/SMS",
        "Configurable rules (schedules + event combinations)",
        "Search and reports for suspicious events",
        "Easy PC operator interface (Milestone Smart Client)",
        "GDPR/AEPD-oriented configuration and documentation",
        "Remote training + 3–6 months post-deployment support",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    # 3. Scope
    add_heading(doc, "3. Project Scope", 1)
    add_heading(doc, "3.1 In Scope", 2)
    for item in [
        "Remote audit of server/network/camera readiness and Milestone licensing status",
        "Milestone XProtect configuration (cameras, views, recording, users/roles)",
        "ONVIF integration/optimization for Safire Smart E1 cameras",
        "AI analytics configuration (facial watchlists + multi line-crossing)",
        "Milestone rules/alarms for schedules and combined events",
        "Alert path setup for App/mobile + WhatsApp/SMS via approved connector",
        "GDPR/AEPD baseline settings: retention, privacy masks, access roles",
        "Remote operator training sessions",
        "Technical documentation + end-user manual",
        "Post-go-live remote support (selected package)",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    add_heading(doc, "3.2 Out of Scope", 2)
    for item in [
        "Physical camera mounting / cabling (already installed)",
        "Purchase of Milestone / BriefCam / WhatsApp Business licenses (client-provided unless agreed otherwise)",
        "On-site travel to Valencia",
        "Development of a custom VMS to replace Milestone",
        "Legal GDPR advisory / DPO services (technical configuration + templates only)",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    # 4. Tech stack
    add_heading(doc, "4. Technology Stack", 1)
    add_table(
        doc,
        ["Layer", "Technology", "Role in project"],
        [
            ["VMS Core", "Milestone XProtect", "Central video management, recording, users, rules/alarms"],
            ["Operator UI", "XProtect Smart Client", "Live view, playback, alarm handling on PC"],
            ["Admin UI", "XProtect Management Client", "System configuration and rule setup"],
            ["Cameras", "Safire Smart E1 6MP (ONVIF)", "Video sources (Bullet/Turret)"],
            ["Edge/NVR", "Safire Smart B2 NVR", "Existing recording appliance / bridge as applicable"],
            ["AI Analytics", "BriefCam or equivalent", "Facial recognition, watchlists, advanced search/events"],
            ["Events/Rules", "Milestone Rules + Alarm Manager", "Schedules, combinations, severity routing"],
            ["Mobile access", "Milestone Mobile / Web (as licensed)", "Remote monitoring for authorized staff"],
            ["WhatsApp/SMS", "WhatsApp Business API / SMS gateway (e.g. Twilio or approved provider)", "Outbound critical alerts"],
            ["Identity/Access", "Milestone roles + Windows/AD if used", "Admin / Security / Reception permissions"],
            ["Compliance", "Privacy masks, retention policies, audit logs", "GDPR/AEPD technical controls"],
            ["Docs & Demo", "Technical manuals + workflow demo", "Training aid and handover package"],
        ],
    )
    add_para(
        doc,
        "Note: Facial recognition and WhatsApp are typically connected to Milestone (AI module + messaging connector), "
        "not native one-click features inside every XProtect edition. Final connector choice is confirmed in Phase 1.",
        size=10,
    )

    # 5. Approach
    add_heading(doc, "5. Delivery Approach", 1)
    for item in [
        "Phase-gated remote delivery with acceptance after each milestone",
        "Secure remote access (VPN/RDP) to Milestone server and camera network",
        "Configure first on a pilot set of cameras, then roll out to all 30",
        "Tune analytics to reduce false alarms before go-live",
        "Provide bilingual-ready operator guidance (ES/EN as needed)",
        "Keep production changes reversible and documented",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    # 6. Milestones + budget
    add_heading(doc, "6. Milestones, Timeline & Budget", 1)
    add_para(
        doc,
        "Currency: EUR. Rate basis aligned with project posting (~€15/hour average). "
        "Total recommended go-live package: €2,100. Optional support packages listed below.",
    )

    add_table(
        doc,
        ["Milestone", "Duration", "Key activities", "Deliverables", "Budget (EUR)"],
        [
            [
                "M1 — Analysis & Kickoff",
                "4–5 days",
                "Access check, license audit, network/camera health, GDPR requirements workshop",
                "Audit report + implementation plan + risk list",
                "250",
            ],
            [
                "M2 — Milestone Core Setup",
                "5–7 days",
                "Camera onboarding (30), views/maps, recording profiles, users & roles",
                "Configured XProtect environment ready for analytics",
                "550",
            ],
            [
                "M3 — AI Analytics & Rules",
                "7–10 days",
                "Watchlists, multi line-crossing, schedules, combined event rules",
                "AI rules matrix + tuned detection profiles",
                "650",
            ],
            [
                "M4 — Alerting (Mobile/WhatsApp/SMS)",
                "3–4 days",
                "Alarm routing, connector setup, recipient matrix, escalation tests",
                "Working alert path with test evidence",
                "300",
            ],
            [
                "M5 — Testing & Optimization",
                "3–4 days",
                "Scenario tests (night access, blocked list, door lines), false-alarm reduction",
                "Test protocol + sign-off checklist",
                "200",
            ],
            [
                "M6 — Training & Documentation",
                "2–3 days",
                "Remote training for hotel staff, manuals, handover",
                "User manual + technical docs + recorded session (optional)",
                "150",
            ],
        ],
    )

    add_para(doc, "Go-live subtotal (M1–M6): €2,100", bold=True)
    add_para(doc, "Estimated active delivery time to go-live: 5–6 weeks", bold=True)

    add_heading(doc, "6.1 Optional Support Packages", 2)
    add_table(
        doc,
        ["Package", "Period", "Includes", "Budget (EUR)"],
        [
            ["Support S3", "3 months", "Remote incident support, rule fine-tuning, monthly health check", "300"],
            ["Support S6", "6 months", "All S3 items + priority response + quarterly optimization review", "550"],
        ],
    )

    add_para(doc, "Recommended commercial package for award:", bold=True)
    add_para(doc, "€2,400 = Go-live (€2,100) + 3 months support (€300)", bold=True)
    add_para(doc, "Alternative: €2,650 = Go-live (€2,100) + 6 months support (€550)")

    # Payment
    add_heading(doc, "7. Payment Schedule", 1)
    add_table(
        doc,
        ["Payment", "When", "Amount"],
        [
            ["P1", "Contract start / after M1 kickoff acceptance", "20% (€420 of go-live)"],
            ["P2", "After M2 + M3 acceptance", "40% (€840)"],
            ["P3", "After M4 + M5 acceptance", "25% (€525)"],
            ["P4", "After M6 training/docs acceptance (go-live)", "15% (€315)"],
            ["Support", "Monthly or upfront for selected support package", "Per package"],
        ],
    )

    # Deliverables
    add_heading(doc, "8. Deliverables Checklist", 1)
    for item in [
        "Milestone XProtect configured for hotel operations",
        "30 cameras integrated and organized by floor/zone",
        "Facial watchlists configured (employees/guests/blocked)",
        "Virtual line-crossing rules for doors and common areas",
        "Alert workflow to mobile + WhatsApp/SMS (as approved)",
        "Role-based operator access (Admin / Security / Reception)",
        "GDPR-oriented technical settings and documentation pack",
        "Remote training completed",
        "Support handover package",
    ]:
        add_para(doc, f"☐ {item}", space_after=2)

    # Assumptions
    add_heading(doc, "9. Assumptions & Client Responsibilities", 1)
    for item in [
        "Client provides remote admin access to Milestone server, NVR, and camera network",
        "Valid Milestone XProtect license/edition is available (or procurement is handled by client)",
        "AI analytics license (BriefCam or equivalent) is available, or approved as an add-on purchase",
        "WhatsApp Business / SMS provider account is available, or connector cost is approved",
        "Client nominates alert recipients and escalation contacts",
        "Client provides retention policy preferences and privacy-mask zones",
        "Network bandwidth and server hardware meet Milestone/AI vendor minimums",
        "Changes to scope (extra sites, extra analytics modules, custom development) may adjust budget/timeline",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    # Risks
    add_heading(doc, "10. Risks & Mitigations", 1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Missing AI/WhatsApp licenses", "Delay on facial/WhatsApp features", "Confirm in M1; provide fallback (SMS/email/mobile app)"],
            ["High false alarms", "Operator fatigue", "Pilot tuning before full rollout; scheduled sensitivity reviews"],
            ["Unstable remote access", "Delivery slip", "Dedicated VPN account + change window calendar"],
            ["Camera ONVIF limitations", "Partial analytics metadata", "Validate per camera model in M1; use camera-side events where needed"],
        ],
    )

    # Demo
    add_heading(doc, "11. Demo Provided", 1)
    add_para(
        doc,
        "An interactive workflow demo is available to show the intended operator experience "
        "(Introduction, Operations, Live view, Watchlists, Rules, Alerts, Search, GDPR, ES/EN). "
        "This demo is a simulation for clarification and training alignment. Production implementation "
        "is performed on Milestone XProtect with the hotel’s Safire infrastructure.",
    )

    # Why us / closing
    add_heading(doc, "12. Why This Plan Works", 1)
    for item in [
        "Matches the client’s required platform (Milestone-first, not a custom VMS rewrite)",
        "Clear milestones with fixed budgets and acceptance points",
        "Fully remote delivery (no Valencia travel cost/time)",
        "Includes compliance-minded configuration and staff enablement",
        "Support options protect the hotel after go-live",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    add_heading(doc, "13. Acceptance", 1)
    add_para(doc, "If this proposal is accepted, kickoff can start upon:")
    for item in [
        "Confirmation of selected package (Go-live only / +3 months / +6 months support)",
        "Remote access credentials",
        "License confirmation for Milestone + AI + messaging",
        "Signed milestone payment terms",
    ]:
        add_para(doc, f"• {item}", space_after=2)

    add_para(doc, "")
    add_para(doc, "Prepared by: IntelliForge", bold=True)
    add_para(doc, "Project: Intelligent Hotel Video Surveillance — Milestone XProtect + AI Analytics")
    add_para(doc, "Total recommended investment: €2,400 (go-live + 3 months support)")
    add_para(doc, "Timeline to go-live: approximately 5–6 weeks (remote)")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
